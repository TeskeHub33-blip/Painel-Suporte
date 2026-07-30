# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x24, 0x24, 0x3E)
NAVY_DARK = RGBColor(0x16, 0x16, 0x28)
GREY_SUB = RGBColor(0x59, 0x59, 0x74)
GREY_BODY = RGBColor(0x40, 0x40, 0x40)
PINK = RGBColor(0xED, 0x6D, 0xA2)
PINK_DARK = RGBColor(0xE0, 0x55, 0x92)

FONT_BODY = "Avenir Next"
FONT_TITLE = "Avenir Next Medium"
LOGO = "C:/Users/WellingtonErvinoTesk/Documents/Claude/manual_suporte/ref_assets/logo_emiteai.png"

doc = Document()
normal = doc.styles['Normal']
normal.font.name = FONT_BODY
normal.font.size = Pt(10.5)

def style_run(run, size=10.5, color=GREY_BODY, bold=False, italic=False, font=FONT_BODY):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)

def add_body(text, size=10, color=GREY_BODY, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    style_run(r, size=size, color=color, bold=bold, italic=italic)
    return p

def add_h1(text):
    doc.add_page_break() if doc.paragraphs and len(doc.paragraphs) > 3 else None
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '6'); bottom.set(qn('w:color'), '24243E')
    border.append(bottom); pPr.append(border)
    r = p.add_run(text)
    style_run(r, size=17, color=NAVY, bold=True, font=FONT_TITLE)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=12.5, color=NAVY_DARK, bold=True)

def add_card(nome, tipo, regra, notas=None):
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    row = table.add_row().cells[0]
    tcPr = row._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F2F2F2')
    tcPr.append(shd)
    p1 = row.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(nome)
    style_run(r1, size=11, color=NAVY_DARK, bold=True)
    r1b = p1.add_run("  [" + tipo + "]")
    style_run(r1b, size=8.5, color=PINK_DARK, bold=True)

    p2 = row.add_paragraph()
    p2.paragraph_format.space_after = Pt(2) if notas else Pt(0)
    r2 = p2.add_run("Regra: ")
    style_run(r2, size=9.5, color=GREY_BODY, bold=True)
    r2b = p2.add_run(regra)
    style_run(r2b, size=9.5, color=GREY_BODY)

    if notas:
        p3 = row.add_paragraph()
        r3 = p3.add_run("Observacao: ")
        style_run(r3, size=9, color=GREY_SUB, bold=True, italic=True)
        r3b = p3.add_run(notas)
        style_run(r3b, size=9, color=GREY_SUB, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------------- CAPA ----------------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(); r.add_picture(LOGO, width=Cm(6.5))
for _ in range(3): doc.add_paragraph()
p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run("Manual de Regras"); style_run(r1, size=26, color=NAVY, bold=True, font=FONT_TITLE)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Painel Diario de Suporte — o que cada card/KPI de cada aba calcula")
style_run(r2, size=12.5, color=GREY_SUB, italic=True)
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Fonte: Movidesk (API). Board publicado em teskehub33-blip.github.io/Painel-Suporte")
style_run(r3, size=9.5, color=GREY_SUB, italic=True)
p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Atualizado em Julho/2026 — cobre as abas Ao Vivo, Historico, Clientes, One-on-One, Gamificacao e Reuniao Mensal")
style_run(r4, size=9, color=GREY_SUB, italic=True)
doc.add_page_break()

# ---------------- INTRO ----------------
add_h1("Como ler este manual")
add_body(
    "Cada card do painel e um filtro/calculo aplicado sobre os dados do Movidesk, recalculado a cada "
    "atualizacao automatica (a cada 15 minutos, via GitHub Actions disparado por um gatilho externo). "
    "Este documento lista, aba por aba, exatamente qual regra esta por tras de cada numero — para que "
    "qualquer pessoa da equipe consiga auditar ou questionar um valor exibido."
)
add_body(
    "Convencao usada abaixo: [ao vivo] = calculado sobre os chamados atualmente ABERTOS (nao fechados/"
    "cancelados/resolvidos); [historico] = calculado sobre chamados RESOLVIDOS (hoje ou nos ultimos 3 "
    "meses, conforme indicado), e sempre restrito ao time de Suporte (ownerTeam = 'Suporte')."
)
add_body(
    "Regra transversal de exclusao (Cancelado + Azure): TODOS os conjuntos de chamados resolvidos "
    "(RESOLVED_MONTH e RESOLVED_MONTHS, usados por Historico, Clientes, One-on-One, Gamificacao e "
    "Reuniao Mensal) ja excluem, na origem, chamados com status = 'Cancelado' e chamados reabertos "
    "INDEVIDAMENTE pela integracao com o Azure (que reabre o chamado logo apos ele ja ter sido marcado "
    "Resolvido/Fechado, sem uma pessoa envolvida). Esse filtro fica em um unico ponto do codigo "
    "(construcao de RESOLVED_MONTH_ALL/RESOLVED_MONTHS) para garantir que TODAS as abas fiquem "
    "automaticamente consistentes — nenhuma aba precisa de um filtro proprio para isso.",
    bold=False
)
add_body(
    "Reaberturas LEGITIMAS (feitas por cliente ou agente) continuam contando normalmente. A deteccao de "
    "reabertura indevida usa uma heuristica sobre o historico de status do chamado: se a reabertura "
    "aconteceu em menos de 60 minutos apos a transicao anterior ter sido Resolvido/Fechado, e tratada "
    "como indevida (tempo curto demais para ser uma acao manual). So e possivel verificar essa "
    "heuristica no mes corrente (onde o historico de status completo esta disponivel) — nos dois meses "
    "anteriores, sem esse historico, todo chamado reaberto e tratado como legitimo por padrao, para nao "
    "penalizar avaliacoes antigas por engano.",
    bold=False
)
add_body(
    "Abas protegidas por senha (One-on-One, Gamificacao, Reuniao Mensal): a senha (3300) e apenas uma "
    "trava simples do lado do navegador para evitar acesso casual — nao e seguranca real (visivel no "
    "codigo-fonte da pagina). Serve para impedir que qualquer pessoa que acesse o link publico veja "
    "dados individuais de performance sem querer.",
    italic=True
)

# ================= ABA AO VIVO =================
add_h1("Aba \"Ao Vivo\" — KPIs principais")

add_card("Novos (aguard. atend.)", "ao vivo",
    "status do chamado = 'Novo'.")

add_card("Em atendimento", "ao vivo",
    "status do chamado = 'Em atendimento'.",
    "Cor do card por faixa: verde quando < 20 chamados, amarelo entre 20 e 50, vermelho acima de 50.")

add_card("Aguardando cliente", "ao vivo",
    "status do chamado = 'Aguardando Cliente'.",
    "Cor do card por faixa: verde quando < 20 chamados, amarelo entre 20 e 50, vermelho acima de 50.")

add_card("Aging (Em atendimento > 2 dias)", "ao vivo",
    "time do chamado = 'Suporte' E status = 'Em atendimento' E tempo desde a ultima atualizacao "
    "(lastUpdate) >= 48 horas.",
    "Renomeado de \"Bouncing\" para \"Aging\" (mesma regra de calculo). Cor do card por faixa: verde "
    "quando 0, amarelo ate 10, vermelho acima de 10.")

add_card("Priorizados (WhatsApp)", "ao vivo",
    "time do chamado = 'Suporte' E o chamado possui a tag 'Priorizado' no Movidesk.",
    "Abaixo do numero, mostra o tempo medio de resolucao no mes comparando chamados priorizados vs. "
    "nao priorizados (mesma tag), para acompanhar se a priorizacao esta de fato acelerando o atendimento. "
    "O conjunto de chamados abertos (TICKETS) ja exclui, na propria consulta a API do Movidesk, os "
    "status Fechado/Cancelado/Resolvido — logo o card nunca soma chamados ja concluidos.")

add_card("Contraturno em atendimento", "ao vivo",
    "status = 'Em atendimento' E tecnico responsavel (owner) e Alife Caetano dos Santos OU Vinicius "
    "Campestrini.",
    "Contraturno e definido pelos TECNICOS do turno, nao por horario de criacao do chamado.")

add_card("Carga parada / CIOT-MDFe-CTe", "ao vivo",
    "status = 'Em atendimento' E (o ASSUNTO do chamado OU a DESCRICAO da primeira acao do chamado) "
    "contem, via expressao regular, termos como 'carga trava/parad', 'travad', 'ciot', 'mdfe'/'mdf-e', "
    "'cte'/'ct-e' (case-insensitive).",
    "Deteccao por palavra-chave no titulo OU no corpo da descricao do chamado (antes so verificava o "
    "titulo) — nao confirma o status real da carga no sistema, apenas indica indicio textual. Cor: "
    "verde quando 0, vermelho quando > 0.")

add_card("Possivel classificacao incorreta", "ao vivo",
    "categoria do chamado esta em {Duvida, Erro Operacional, Terceiros} E o chamado ja passou, em algum "
    "momento do seu historico de status, pelo status 'Aguardando Desenvolvimento - fila Bugs'.",
    "Logica: Melhoria, Bug e (alguns) Servicos legitimamente geram task/fila de dev; Duvida/Erro "
    "Operacional/Terceiros nao deveriam — se passaram pela fila de Bugs mesmo assim, e sinal de "
    "categoria errada. Cor: verde quando 0, amarelo quando > 0.")

add_h1("Aba \"Ao Vivo\" — Fila de Priorizacao Operacional")
add_body(
    "Reordena os chamados abertos em 2 niveis de prioridade (niveis 3 \"Recorrencia/melhoria\" e 4 "
    "\"Outros\" foram REMOVIDOS a pedido — a fila agora so lista chamados que se encaixam em bloqueio "
    "operacional ou risco fiscal; chamados que nao se enquadram em nenhum dos dois niveis simplesmente "
    "nao aparecem na fila). Dentro de cada nivel, os chamados mais antigos (maior tempo aberto) aparecem "
    "primeiro."
)
add_card("Nivel 1 — Bloqueio operacional", "prioridade",
    "assunto contem (regex) mdfe/mdf-e, ciot, gnre, integracao/integracoes, ou carga travada/parada — "
    "E o assunto NAO contem termos de risco fiscal (ver Nivel 2).")
add_card("Nivel 2 — Risco fiscal (multas)", "prioridade",
    "assunto contem (regex) multa, risco fiscal, imposto, difal, icms, ou mencao a guia vencida/vencimento "
    "de guia.")

add_h1("Aba \"Ao Vivo\" — paineis com listas/graficos")
add_card("Chats em atendimento — quem e ha quanto tempo", "ao vivo",
    "status = 'Em atendimento' E (origin do chamado = 24 OU chatGroup preenchido). Lista tecnico e tempo "
    "desde a ultima atualizacao.",
    "APROXIMACAO, ainda pendente de ajuste: o ideal e monitorar a FILA de chat (sessoes de chat em "
    "andamento), nao os chamados de origem chat ja abertos no Movidesk — mas o Movidesk nao expõe uma "
    "API publica de fila de chat em tempo real, entao o card continua usando origin=24/chatGroup como "
    "indicio de que o chamado se originou por canal de chat, sem garantir que existe uma sessao "
    "acontecendo neste momento.")
add_card("Em atendimento por tecnico", "ao vivo",
    "agrupa o card \"Em atendimento\" (status = 'Em atendimento') por tecnico responsavel (owner), "
    "ordenado do tecnico com mais chamados para o com menos.")
add_card("Nao atualizados hoje por tecnico / Nao atualizados hoje (lista)", "ao vivo",
    "status em {Em atendimento, Aguardando Cliente} E a data da ultima atualizacao (lastUpdate) e "
    "diferente da data de hoje (comparando apenas a parte de data, em UTC).")

# ================= ABA HISTORICO =================
add_h1("Aba \"Historico\" — escopo geral")
add_body(
    "IMPORTANTE: todos os calculos desta aba partem de conjuntos de dados ja filtrados na origem "
    "para conter apenas chamados do time 'Suporte' (ownerTeam = 'Suporte') — chamados de Implantacao, "
    "Customer Success etc. nunca entram nestes numeros:"
)
add_body("• Resolvidos hoje: chamados com resolvedIn a partir de 00:00 de hoje (UTC).", size=9.5)
add_body("• Resolvidos no mes: chamados com resolvedIn dentro do mes corrente (do dia 1 ate agora).", size=9.5)
add_body(
    "• Filtro por cliente (seletor no canto superior direito da aba): quando um cliente e selecionado, "
    "TODOS os cards abaixo do mes corrente (SLA, MTTR, 1a resposta, bugs, resolvidos por tecnico, e os "
    "drill-downs/exportacoes) passam a considerar somente os chamados daquele cliente. A media de 3 "
    "meses tambem e recalculada so para aquele cliente. \"Todos os clientes\" volta ao calculo geral.",
    size=9.5, space_after=10
)

add_h1("Aba \"Historico\" — KPIs do topo")
add_card("Resolvidos c/ 1a resposta (hoje) / (mes)", "historico",
    "dentro do conjunto de resolvidos (hoje ou mes), conta quantos tem actionCount <= 3 (abertura do "
    "cliente + retorno automatico + a resposta que ja resolveu — no maximo 2 respostas apos a abertura).",
    "Substitui o campo nativo 'resolvedInFirstCall' do Movidesk, que usa outro criterio interno. O card "
    "do mes ainda mostra a media dos ultimos 3 meses e uma meta dinamica: +10% sobre essa media.")
add_card("SLA atendido no prazo (mes)", "historico",
    "dentro dos chamados resolvidos no mes que possuem slaSolutionDate definido E nao foram reabertos indevidamente (ver regra transversal de SLA) "
    "(reopenedIn vazio), conta quantos foram resolvidos (resolvedIn) ANTES OU NA data-limite do SLA "
    "(resolvedIn <= slaSolutionDate). Exibido como percentual do total.",
    "Mostra tambem a media dos ultimos 3 meses e uma meta dinamica de +10% sobre essa media.")
add_card("Tempo medio de atendimento (MTTR)", "historico",
    "para os chamados resolvidos no mes corrente, EXCLUINDO categoria = 'Melhoria', com createdDate e "
    "resolvedIn validos, calcula a media do tempo entre abertura e resolucao.",
    "Melhorias tem ciclo de desenvolvimento proprio (fila de dev) e distorciam a media para cima — por "
    "isso passaram a ser desconsideradas deste indicador. Mostra o detalhamento mes a mes dos ultimos 3 "
    "meses (tambem sem Melhoria), a media desses 3 meses, e uma meta dinamica: 10% MENOR que essa media "
    "(menor e melhor para tempo). Card fica verde quando a meta e batida.")
add_card("Chats resolvidos (mes)", "historico",
    "dentro dos resolvidos no mes, conta os que tem origin = 24 (chamados originados por canal de chat).",
    "Ao contrario do painel \"Chats em atendimento\" da aba Ao Vivo, este e um numero HISTORICO real "
    "(chamado ja resolvido), sem a limitacao de nao saber se e uma sessao ativa.")

add_h1("Aba \"Historico\" — Ciclo de vida do Bug (Media / Alta)")
add_body(
    "Calculado apenas sobre chamados com categoria = 'Bug', resolvidos no mes corrente, com historico de "
    "status disponivel — e sempre dividido em duas linhas separadas: urgencia 'Media' e urgencia 'Alta'. "
    "So funciona para o mes corrente porque o historico de status (statusHistories) nao e mantido para "
    "meses anteriores, por questao de tamanho dos dados."
)
add_card("Tempo medio para abrir bug", "historico · bug",
    "para cada bug, soma o TEMPO UTIL (permanencyTimeWorkingTime, que ja exclui horario fora do "
    "expediente) de cada etapa do historico de status ATE a primeira vez que o chamado entrou no status "
    "'Aguardando Desenvolvimento - fila Bugs', desconsiderando o tempo em que o chamado ficou em "
    "'Aguardando Cliente' (tempo que nao e responsabilidade do suporte). Media sobre os bugs que tem "
    "essa transicao registrada.",
    "Antes usava tempo corrido (createdDate ate a entrada na fila); passou a usar tempo util em fila, "
    "excluindo Aguardando Cliente, para refletir so o tempo realmente gasto pelo suporte. Clicar no card "
    "abre a lista dos chamados que entram nessa media (drill-down). Sem meta/limite definido ainda — "
    "card em cor neutra.")
add_card("Tempo medio aberto no devops", "historico · bug",
    "soma, para cada bug, o tempo (permanencyTimeFullTime) em TODAS as passagens pelo status "
    "'Aguardando Desenvolvimento - fila Bugs' (um bug pode entrar e sair da fila mais de uma vez). "
    "Media sobre os bugs que passaram pela fila ao menos uma vez.",
    "Clicar no card abre a lista dos chamados que entram nessa media (drill-down).")
add_card("Tempo medio em validacao", "historico · bug",
    "para cada bug, soma o tempo em status 'Em atendimento' ou 'Aguardando Cliente' APOS a ultima saida "
    "da fila de Bugs, ate o chamado ser resolvido. Representa o tempo que o suporte leva pra validar a "
    "correcao com o cliente depois que o dev devolveu o card.",
    "Clicar no card abre a lista dos chamados que entram nessa media (drill-down).")

add_h1("Aba \"Historico\" — SLA por categoria e ranking por tecnico")
add_card("SLA por categoria (tabela + Total geral)", "historico",
    "mesma logica do KPI \"SLA atendido no prazo\" (exclui so as reaberturas indevidas), quebrada por categoria "
    "do chamado (Bug, Duvida, Melhoria etc.), com uma linha de Total geral somando todas as categorias.")
add_card("Chats resolvidos por tecnico (mes)", "historico",
    "agrupa os \"Chats resolvidos (mes)\" (origin = 24) por tecnico responsavel, ordenado do maior "
    "volume para o menor.")
add_card("Chamados resolvidos por tecnico (mes)", "historico",
    "agrupa TODOS os chamados resolvidos no mes (qualquer categoria/origem) por tecnico responsavel. "
    "Mostra tambem a media dos ultimos 3 meses no total geral.")

add_h1("Aba \"Historico\" — Situacoes recorrentes no mes")
add_card("Situacoes recorrentes no mes", "historico",
    "agrupa os chamados resolvidos no mes pelo ASSUNTO normalizado (minusculas, sem acentos/pontuacao, "
    "espacos colapsados) e lista apenas os grupos com 2 ou mais chamados, ordenados do mais recorrente "
    "para o menos. Para cada grupo, mostra a quantidade de ocorrencias e os clientes 'ofensores' "
    "(clientOrg de cada chamado do grupo).",
    "Card novo — respondendo ao pedido de identificar situacoes que se repetiram no mes e apontar quem "
    "sao os ofensores. Clicar em um grupo abre a lista dos chamados daquele grupo (drill-down).")

# ================= ABA CLIENTES =================
add_h1("Aba \"Clientes\" — Status Report por cliente")
add_body(
    "Seletores no canto superior direito: mes (dos ultimos 3 disponiveis) e cliente. O cliente e "
    "identificado a partir do campo 'clients' do Movidesk — prioriza o contato do tipo organizacao "
    "(pessoa juridica); se so houver contato pessoa fisica, tenta inferir a empresa pelo dominio do "
    "e-mail (ignorando provedores genericos tipo gmail/hotmail); sem nenhuma das duas informacoes, "
    "aparece como 'Sem cliente' — nunca usa o nome de uma pessoa fisica como se fosse o cliente."
)
add_card("Chamados no mes", "clientes",
    "total de chamados resolvidos no mes selecionado para o cliente selecionado (time Suporte).",
    "Mostra tambem a media dos ultimos 3 meses do mesmo cliente.")
add_card("MTTR (tempo medio resolucao)", "clientes",
    "media do tempo entre createdDate e resolvedIn, dos chamados do cliente no mes selecionado.")
add_card("SLA no prazo", "clientes",
    "% de chamados do cliente, no mes selecionado, com slaSolutionDate definido E sem reabertura indevida, "
    "resolvidos dentro do prazo. Cor vermelha quando abaixo de 70%.")
add_card("Backlog (em aberto agora)", "clientes",
    "quantidade de chamados do cliente ATUALMENTE abertos (nao fechados/cancelados/resolvidos) — "
    "e um numero AO VIVO, independente do mes selecionado no filtro.")
add_card("Reincidencia (reabertos)", "clientes",
    "quantidade e percentual de chamados do cliente, no mes selecionado, que possuem o campo "
    "reopenedIn preenchido (foram reabertos ao menos uma vez).")
add_card("Gestao de Chamados — backlog em aberto agora", "clientes",
    "tabela com o backlog atual do cliente (mesmos chamados do card \"Backlog\"), agrupado por "
    "categoria e depois por status dentro de cada categoria — no mesmo formato do status report que o "
    "CS apresenta ao cliente. Clicavel: abre a lista de chamados daquela categoria/status.")

# ================= ABA ONE-ON-ONE =================
add_h1("Aba \"One-on-One\" — indicadores individuais (protegida por senha)")
add_body(
    "Seletores: periodo (um dos ultimos 3 meses) e tecnico responsavel. Ao selecionar o tecnico, um "
    "selo indica o nivel dele: N1 ou N2 — lista fixa configurada no codigo (N2_TECNICOS). Os KPIs e "
    "metas mudam de acordo com o nivel."
)
add_card("Chamados resolvidos", "one-on-one",
    "total de chamados resolvidos pelo tecnico no periodo selecionado.",
    "Mostra a media da equipe do MESMO nivel (N1 ou N2) no periodo, e o detalhamento mes a mes dos "
    "ultimos 3 meses do proprio tecnico.")
add_card("Resolvidos na 1a resposta", "one-on-one",
    "% de chamados do tecnico, no periodo, com actionCount <= 3.",
    "Mostra a media da equipe do mesmo nivel e a media dos ultimos 3 meses do tecnico.")
add_card("Tempo medio de resolucao / SLA no prazo", "one-on-one",
    "mesma logica do MTTR e do SLA da aba Clientes (exclui so reaberturas indevidas), mas escopados ao tecnico e "
    "periodo selecionados.",
    "Comparado contra a media da equipe do mesmo nivel no mesmo periodo.")
add_card("Aging atual (>2 dias) / Nao atualizados hoje", "one-on-one",
    "mesma regra dos cards equivalentes da aba Ao Vivo (renomeado de \"Bouncing\" para \"Aging\"), mas "
    "filtrados pelos chamados atualmente atribuidos ao tecnico selecionado.",
    "Sao numeros AO VIVO — nao mudam com o periodo selecionado. Cor por faixa: verde quando 0, amarelo "
    "ate 5, vermelho acima de 5 (antes era binario verde/vermelho).")
add_card("Metas: 0 aging / 0 nao atualizados", "one-on-one · meta",
    "meta fixa: aging = 0 e nao-atualizados = 0. Card mostra 'Meta batida' (verde) quando o valor "
    "atual e 0, ou 'Meta nao batida' caso contrario.",
    "Sao as unicas metas com valor fixo — os demais indicadores usam meta dinamica (ver abaixo) "
    "porque aging/nao-atualizados sao indicadores ao vivo, sem serie mensal para calcular uma base.")
add_card("Meta MTTR / Meta SLA / Meta 1a resposta", "one-on-one · meta",
    "meta dinamica: 10% de melhoria sobre a PROPRIA media dos ultimos 3 meses do tecnico. Para MTTR "
    "(tempo), a meta e a media * 0.9 (10% menor). Para SLA e 1a resposta (percentuais de qualidade), "
    "a meta e a media * 1.1 (10% maior, limitado a 100%).",
    "Card mostra 'Meta batida'/'Meta nao batida' comparando o valor do periodo selecionado contra essa "
    "meta calculada.")
add_card("Indicadores tecnicos N2 (Bug/Melhoria/Servicos)", "one-on-one · N2",
    "somente visivel quando o tecnico selecionado e N2: total de chamados tecnicos (categorias Bug, "
    "Melhoria, Servicos) resolvidos, % desses que passaram pela fila de Bugs (task associada), e tempo "
    "medio em devops/validacao — mesma logica do ciclo de vida do Bug, mas por tecnico.",
    "So calculavel para o mes corrente (offset 0), pois depende do historico de status "
    "(statusHistories), que nao e mantido para meses anteriores.")
add_card("Indicadores tecnicos N1 — tempo ate acionar N2", "one-on-one · N1",
    "somente visivel quando o tecnico selecionado e N1: total de chamados tecnicos (categorias Bug, "
    "Melhoria, Servicos) atribuidos ao tecnico no mes, quantos desses chegaram a entrar na fila de "
    "Bugs (acionamento do N2), e o tempo medio (util, excluindo Aguardando Cliente) desde a atribuicao "
    "do chamado ao tecnico ate a entrada na fila de Bugs.",
    "APROXIMACAO: o Movidesk nao mantem um historico de troca de responsavel (owner), apenas de status "
    "— entao nao existe um timestamp exato de \"quando o N1 acionou o N2\". Este indicador usa a entrada "
    "na fila de Bugs ('Aguardando Desenvolvimento - fila Bugs') como proxy razoavel para esse momento, ja "
    "que e o evento que efetivamente traz o time de dev/N2 para o chamado. So calculavel para o mes "
    "corrente, pela mesma limitacao de historico de status das metricas N2.")

# ================= ABA GAMIFICACAO =================
add_h1("Aba \"Gamificacao\" — metas batidas por criterio (protegida por senha)")
add_body(
    "Nao mostra nomes de tecnicos individuais — agrega quantas vezes cada CRITERIO foi batido, "
    "somando todos os tecnicos do periodo selecionado (ou a soma dos ultimos 3 meses)."
)
add_card("Metas batidas (soma geral)", "gamificacao",
    "para cada tecnico e cada um dos 3 criterios (MTTR, SLA no prazo, 1a resposta), verifica se a meta "
    "de 10% de melhoria sobre a propria media de 3 meses do tecnico foi batida naquele periodo. Soma "
    "quantas avaliacoes tecnico x criterio bateram a meta, sobre o total de avaliacoes possiveis.",
    "Mesma formula de meta usada na aba One-on-One, aplicada a todos os tecnicos e agregada.")
add_card("Criterios avaliados (tabela)", "gamificacao",
    "tabela com uma linha por criterio (MTTR, SLA no prazo, 1a resposta), mostrando quantas avaliacoes "
    "tecnico x mes bateram a meta, o total avaliado, e o percentual de aproveitamento.",
    "Os criterios de equipe (media da equipe usada como referencia) usam exatamente a mesma formula e a "
    "mesma base de dados dos criterios individuais — nao ha uma logica separada para equipe vs. "
    "individuo.")

# ================= ABA REUNIAO MENSAL =================
add_h1("Aba \"Reuniao Mensal\" — fechados por categoria + metas (protegida por senha)")
add_card("SLA no prazo (KPI por mes)", "reuniao mensal",
    "mesma logica do SLA da aba Historico (exclui so reaberturas indevidas), um card por mes dos ultimos 3 meses "
    "disponiveis.")
add_card("Tabela: Chamados fechados por categoria", "reuniao mensal",
    "uma linha por mes, com a contagem de chamados fechados em cada categoria fixa (Bloqueio Sistema, "
    "Bug, Duvida, Melhoria, Erro Operacional, Terceiros, Servicos, GNRE Pagamento — mais uma coluna "
    "'Outros' para qualquer categoria fora dessa lista), o Total Fechados do mes, o % de Duvidas sobre "
    "o total, a Reducao vs. meta de duvidas, e o % SLA no prazo do mes. A ultima linha soma tudo no "
    "periodo (Total).")
add_card("Meta de duvidas: reduzir 20% (baseline 61%)", "reuniao mensal · meta",
    "baseline fixo de 61% (media historica de referencia, calculada sobre ~7 meses e 3535 chamados de "
    "Duvida). A meta e reduzir 20% RELATIVOS sobre esse baseline (nao 20 pontos percentuais). Formula: "
    "reducao% = (61 - %Duvidas_atual) / 61 * 100, arredondado.",
    "Card e coluna da tabela ficam verdes quando a reducao alcancada e >= 20% (meta batida), amarelo "
    "quando ha reducao mas ainda abaixo de 20%, e vermelho quando o % de duvidas piorou em relacao ao "
    "baseline (reducao negativa).")

# ================= NOTAS TECNICAS =================
add_h1("Notas tecnicas gerais")
add_body("• Ciclo de atualizacao: cada regra e recalculada do zero a cada atualizacao automatica (a cada 15 minutos, via GitHub Actions disparado por um cron externo); nao ha acumulo/cache de dados entre ciclos.", size=9.5)
add_body("• Protocolo exibido: todas as tabelas mostram o campo \"protocol\" do Movidesk (nao o id interno); clicar nele abre o chamado real em https://emiteai.movidesk.com/Ticket/Edit/{id}.", size=9.5)
add_body("• \"Tecnico\" = campo owner.businessName do Movidesk (responsavel atual pelo chamado).", size=9.5)
add_body("• Horario exibido no cabecalho do painel esta em horario de Brasilia (UTC-3); os dados internos (createdDate, lastUpdate, resolvedIn) sao processados em UTC.", size=9.5)
add_body("• Exportar para Excel: gera um arquivo .txt separado por tabulacao — abrir no Excel via Dados > Obter Dados > De Texto, delimitador Tabulacao.", size=9.5)
add_body("• Titulos dos cards sao editaveis: duplo clique em qualquer titulo de KPI para renomea-lo; a alteracao fica salva no navegador local (localStorage) e sobrevive as atualizacoes automaticas de dados, mas e por navegador/maquina, nao sincroniza entre pessoas.", size=9.5)
add_body("• Publicacao: o painel e um site estatico publicado no GitHub Pages; os dados sao buscados e o HTML e regenerado por um workflow do GitHub Actions, disparado por um servico externo (cron-job.org) a cada 15 minutos — nao depende de nenhuma maquina local ligada.", size=9.5)

import sys
out = sys.argv[1] if len(sys.argv) > 1 else "C:/Users/WellingtonErvinoTesk/Documents/Claude/manual_suporte/Manual_Regras_Painel_Diario.docx"
doc.save(out)
print("Salvo em:", out)
