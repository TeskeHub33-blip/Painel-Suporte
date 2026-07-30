# -*- coding: utf-8 -*-
import os
import copy
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

BASE = "C:/Users/WellingtonErvinoTesk/Documents/Claude/manual_suporte"
LOGO = f"{BASE}/ref_assets/logo_emiteai.png"
OUT_DIR = f"{BASE}/manuais_emiteai"
os.makedirs(OUT_DIR, exist_ok=True)

# ---- paleta EmiteAi extraida do manual de referencia (SPED Fiscal) ----
NAVY = RGBColor(0x1F, 0x4E, 0x79)      # titulos / tabelas
NAVY_DARK = RGBColor(0x0F, 0x47, 0x61)
GREY_SUB = RGBColor(0x59, 0x59, 0x59)  # subtitulo
GREY_BODY = RGBColor(0x40, 0x40, 0x40) # corpo de texto
PINK = RGBColor(0xEE, 0x6A, 0xA2)      # acento de marca (logo)
PINK_DARK = RGBColor(0xD9, 0x55, 0xC0)
LIGHT_BLUE = "DEEAF1"
LIGHT_GREY = "F2F2F2"
WARN_BG = "FCE4D6"

FONT_BODY = "Avenir Next"
FONT_TITLE = "Avenir Next Medium"

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_page_number_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE   \\* MERGEFORMAT'
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep); r.append(fld_end)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY_SUB

def add_footer_text(section, text):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY_SUB
    run.font.name = FONT_BODY
    run.italic = True

def style_run(run, size=10.5, color=GREY_BODY, bold=False, italic=False, font=FONT_BODY):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)

def add_body_paragraph(doc, text, size=10.5, color=GREY_BODY, bold=False, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    style_run(run, size=size, color=color, bold=bold, italic=italic)
    return p

def add_section_title(doc, text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '1F4E79')
    border.append(bottom); pPr.append(border)
    run = p.add_run(text)
    style_run(run, size=14, color=NAVY, bold=True, font=FONT_TITLE)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=11.5, color=NAVY_DARK, bold=True)
    return p

def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(it)
        style_run(run, size=10.5, color=GREY_BODY)

def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(it)
        style_run(run, size=10.5, color=GREY_BODY)

def add_callout(doc, icon, label, text, bg_hex):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{icon} {label}: ")
    style_run(run, size=10.5, color=NAVY_DARK, bold=True)
    run2 = p.add_run(text)
    style_run(run2, size=10.5, color=GREY_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_metric_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = ''
    hdr[1].text = ''
    for i, cell in enumerate(hdr):
        set_cell_shading(cell, '1F4E79')
        p = cell.paragraphs[0]
        run = p.add_run('Indicador' if i == 0 else 'Valor (julho/2026)')
        style_run(run, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    for idx, (label, value) in enumerate(rows):
        row = table.add_row().cells
        row[0].text = ''
        row[1].text = ''
        shade = LIGHT_GREY if idx % 2 == 0 else 'FFFFFF'
        for cell, txt, bold in [(row[0], label, True), (row[1], value, False)]:
            set_cell_shading(cell, shade)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            style_run(run, size=10, color=GREY_BODY, bold=bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def new_document():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    add_page_number_header(section)
    add_footer_text(section, "Manual de Suporte | EmiteAi — Uso interno")
    return doc, section

def add_cover(doc, titulo, subtitulo, protocolo_ref=None):
    # logo centralizado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO, width=Cm(6.5))
    for _ in range(3):
        doc.add_paragraph()

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(titulo)
    style_run(r1, size=24, color=NAVY, bold=True, font=FONT_TITLE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    r2 = p2.add_run(subtitulo)
    style_run(r2, size=12, color=GREY_SUB, italic=True)

    if protocolo_ref:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(f"⚠ Referencia real do periodo: {protocolo_ref}")
        style_run(r3, size=10, color=NAVY_DARK, bold=True)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Manual gerado a partir da analise dos chamados Movidesk de julho/2026 e da base de conhecimento tecnica do EmiteAi.")
    style_run(r4, size=9.5, color=GREY_SUB, italic=True)

    doc.add_page_break()

def add_topic_body(doc, sec):
    add_section_title(doc, "1. Descricao do problema")
    add_body_paragraph(doc, sec['descricao'])

    add_section_title(doc, "2. Indicadores do periodo analisado")
    add_metric_table(doc, sec['metricas'])

    add_section_title(doc, "3. Sintomas tipicos")
    add_bullets(doc, sec['sintomas'])

    add_section_title(doc, "4. Causa raiz")
    add_callout(doc, "\U0001F50D", "Causa raiz", sec['causa_raiz'], LIGHT_BLUE)

    add_section_title(doc, "5. Diagnostico passo a passo")
    add_numbered(doc, sec['diagnostico'])

    add_section_title(doc, "6. Solucao / encaminhamento")
    add_bullets(doc, sec['solucao'])

    add_section_title(doc, "7. Prevencao e monitoria preventiva")
    add_callout(doc, "\U0001F4A1", "Sugestao de prevencao", "", LIGHT_GREY)
    add_bullets(doc, sec['prevencao'])

    if sec.get('confianca'):
        add_body_paragraph(doc, f"Nivel de confianca desta secao: {sec['confianca']}", size=9, color=GREY_SUB, italic=True, space_after=2)

# ============================================================
# CONTEUDO DOS 10 TOPICOS
# ============================================================
topics = []

topics.append({
 'slug': '01_reprocessar_vpo',
 'titulo': 'Reprocessar VPO (Vale-Pedagio)',
 'subtitulo': 'Vale-pedagio travado, rejeitado ou nao emitido apos CIOT autorizado',
 'ref': "cluster 'Reprocessar vpo' - 70+ chamados em julho/2026",
 'descricao': (
    "E o assunto de maior volume no mes: variacoes de 'Reprocessar vpo', 'Vpo rejeitado', "
    "'Erro no VPO' somam mais de 70 chamados. Na maioria dos casos o vale-pedagio da viagem "
    "nao e comprado/emitido automaticamente e o operador de suporte precisa acionar reprocessamento manual "
    "pelo Event Control."
 ),
 'metricas': [
    ('Total de chamados (VPO/pedagio)', '176'),
    ('Tempo medio de resolucao', '19,6 horas'),
    ('Chamados com mais de 48h em aberto', '26'),
    ('Maior recorrencia de assunto identico', "'Reprocessar vpo' (42x + 28x variantes)"),
 ],
 'sintomas': [
    "Vale-pedagio nao aparece na viagem mesmo com CIOT autorizado.",
    "Chamado com titulo generico 'Reprocessar vpo' / 'Vpo reprocessamento' / 'Erro no VPO'.",
    "Vale-pedagio com valor incorreto.",
    "Cluster MELI proprio: 'MELI - REPROCESSAMENTO DE VPO' e 'ERRO VPO - MELI' (14 + 7 ocorrencias).",
 ],
 'causa_raiz': (
    "O VPO depende do CIOT estar autorizado — a compra e disparada pelo evento CIOT_AUTORIZADO, e o identificador "
    "do pedagio e gravado no payload do CIOT sem evento proprio (sem outbox transacional CIOT->VPO). Isso cria "
    "uma janela de corrida onde o consumer do VPO pode nao encontrar o identificador a tempo, resultando em VPO "
    "orfao. Tambem existe o caminho 'VPO sem CIOT' (via REPOM VPO) com regras proprias."
 ),
 'diagnostico': [
    "Confirmar o status do CIOT da viagem (deve estar AUTORIZADO).",
    "Verificar no Event Control / Ticket a timeline do documento (ticket_uuid) buscando o evento de compra de VPO.",
    "Checar se e um caso 'VPO orfao' (CIOT autorizado mas identificador nao propagado a tempo).",
    "Para casos MELI: verificar lastro (ticket_carga_emiteai) antes de reprocessar.",
    "Se o CIOT esta rejeitado/cancelado, nao reprocessar o VPO sem antes resolver o CIOT.",
 ],
 'solucao': [
    "Reprocessar o evento de compra de VPO pelo Event Control apos confirmar CIOT autorizado.",
    "Se for 'VPO orfao' recorrente, escalar para o time de dominio vale-pedagio com ticket_uuid + timestamp.",
    "Tratar o cluster MELI como fila separada — pode exigir fix de integracao, nao so reprocessamento.",
 ],
 'prevencao': [
    "Alerta automatico quando um CIOT fica AUTORIZADO ha mais de X minutos sem VPO associado.",
    "Checklist obrigatorio de suporte: status do CIOT antes de qualquer reprocessamento de VPO.",
    "Levar ao backlog de produto a criacao do outbox transacional CIOT->VPO.",
    "FAQ dedicado ao cluster MELI-VPO.",
 ],
 'confianca': 'Alta para o mecanismo geral (KB do dominio vale-pedagio); media para o cluster MELI especifico.',
})

topics.append({
 'slug': '02_adiantamento_ciot',
 'titulo': 'Adiantamento do CIOT nao reflete no EmiteAi',
 'subtitulo': 'Percentual de adiantamento negociado com o provedor (REPOM) nao aparece na plataforma',
 'ref': 'Chamado #29077 (Bug, Media urgencia, aberto ha 427h no periodo analisado)',
 'descricao': (
    "Cliente informa/gera o CIOT com percentual de adiantamento acordado com o provedor, mas o valor nao aparece "
    "no EmiteAi. Caso real: chamado #29077 'REPOM - Adiantamento no CIOT - Esta refletindo na Repom e nao na "
    "Emiteai em componentes + CIOTS em contigencia nao retornam para Emiteai'."
 ),
 'metricas': [
    ('Chamados no cluster CIOT', '139'),
    ('Tempo medio de resolucao (CIOT)', '47,5 horas'),
    ('Chamados de CIOT com mais de 48h em aberto', '39'),
    ('Chamado de referencia', '#29077'),
 ],
 'sintomas': [
    "Percentual de adiantamento negociado com o provedor nao aparece em nenhuma tela/relatorio do EmiteAi.",
    "'VALOR CIOT VINDO COM O VALOR ERRADO' (alta urgencia, 413h em aberto).",
    "CIOTs em contingencia (codigo verificador 'XXXX') nao retornam/atualizam no EmiteAi.",
 ],
 'causa_raiz': (
    "Existem duas telas para tratar CIOT manualmente: 'Gestao de terceiros > + Novo' (emissao manual, TEM o "
    "campo '% de adiantamento', grava em ciot.valor_adiantamento) e 'Informar CIOT' (apenas vincula um numero "
    "ja existente, SEM campo de adiantamento). Quando o cliente gera o CIOT direto no provedor e so informa o "
    "numero, o valor nunca e capturado. A parametrizacao de CIOT tambem tem 'percentual_adiantamento_fornecedor' "
    "que, se nulo, impede o calculo automatico."
 ),
 'diagnostico': [
    "Identificar qual fluxo o cliente usa: emissao manual (+Novo) ou 'Informar CIOT'.",
    "Consultar parametrizacao_ciot do tenant e verificar percentual_adiantamento_fornecedor (nulo = nao configurado).",
    "Verificar 'criar_conta_pagar_geracao_ciot' e 'gerar_lancamento_adiantamento_separado'.",
    "Para contingencia: 'XXXX' e sentinela esperado — nao e erro; o valor real chega apos nova consulta ao provedor.",
 ],
 'solucao': [
    "Se usa 'Informar CIOT': orientar migrar para emissao manual (+Novo) OU configurar o percentual na parametrizacao.",
    "Ativar 'criar_conta_pagar_geracao_ciot' e 'gerar_lancamento_adiantamento_separado' se esperado ver lancamento financeiro.",
    "Se preso em contingencia, tratar como o manual 'CIOT preso em processando'.",
 ],
 'prevencao': [
    "Alerta na tela de parametrizacao de CIOT quando percentual_adiantamento_fornecedor estiver nulo.",
    "FAQ: 'Por que o adiantamento do meu CIOT nao aparece no EmiteAi?'.",
    "Monitorar tenants que usam majoritariamente 'Informar CIOT' + reclamam de adiantamento.",
 ],
 'confianca': 'Alta — confirmado via consulta direta as tabelas ciot/parametrizacao_ciot em producao e cruzado com o #29077.',
})

topics.append({
 'slug': '03_ciot_rejeitado_cancelamento',
 'titulo': 'CIOT rejeitado, cancelamento e encerramento manual',
 'subtitulo': 'Quando cancelar/encerrar de fato e quando NAO fazer isso',
 'ref': "cluster 'Cte/Ciot rejeitado' - 21+ ocorrencias no mes",
 'descricao': (
    "Chamados pedindo cancelamento, encerramento ou tratativa de rejeicao de CIOT. Inclui 'Cte/Ciot rejeitado' "
    "(21 ocorrencias), 'cancelar ciot' (4x) e pedidos pontuais ('JRPS - Encerrar CIOTs', 'Encerramentos CIOTs')."
 ),
 'metricas': [
    ('Chamados de rejeicao/cancelamento (subset CIOT)', '~25-30'),
    ('Tempo medio do cluster CIOT', '47,5 horas'),
    ('Chamados presos em fila de bugs', '11'),
 ],
 'sintomas': [
    "CIOT aparece REJEITADO mesmo com viagem em transito no provedor.",
    "Pedido manual de 'Cancelar CIOT' / 'Encerrar CIOT' (ex.: 'Encerrar CIOT - 520010806545').",
    "Cliente pede reemissao apos rejeicao repetida.",
 ],
 'causa_raiz': (
    "'Ja cadastrado' (ErrorCode 35 na REPOM) e resposta de idempotencia, nao rejeicao real — tratar como "
    "rejeicao cancela/reemite um CIOT valido indevidamente. Rejeicao cadastral verdadeira (placa/veiculo/"
    "transportador ausente) e outro caminho, corrigido no #17960 para status INCONSISTENTE."
 ),
 'diagnostico': [
    "Antes de cancelar/encerrar, checar no provedor se o CIOT esta realmente ativo.",
    "Se rejeicao e por cadastro ausente, regularizar o cadastro primeiro — reprocessar nao resolve.",
    "Para 'ja cadastrado', confiar no registro local autorizado, sem gerar novo CIOT.",
 ],
 'solucao': [
    "Checklist: 1) checar estado real no provedor, 2) so entao decidir cancelar/reemitir.",
    "Encaminhar para dev quando cadastro ja esta correto e CIOT segue rejeitado sem motivo aparente.",
 ],
 'prevencao': [
    "Treinamento do suporte sobre idempotencia do CIOT (ErrorCode 35).",
    "Campo obrigatorio no Movidesk (checklist) antes de fechar chamados desse tipo.",
 ],
 'confianca': 'Alta — casos e correcoes documentados na KB (dominio ciot).',
})

topics.append({
 'slug': '04_marketplaces',
 'titulo': 'Integracoes de marketplace (Amazon / Shopee / Meli / Magalu)',
 'subtitulo': 'Cargas travadas, eixos, VRID e geracao de documentos nas integracoes',
 'ref': "120 chamados no mes, media de 43,1h para resolucao",
 'descricao': (
    "120 chamados no mes com tempo medio de 43,1h. O subconjunto mais recorrente e VPO/CIOT no fluxo Meli "
    "(ver manual 01), mas ha tambem problemas estruturais de carga (eixos, integracao Shopee, cargas travadas)."
 ),
 'metricas': [
    ('Total de chamados marketplace', '120'),
    ('Tempo medio de resolucao', '43,1 horas'),
    ('Chamados com mais de 48h em aberto', '35'),
 ],
 'sintomas': [
    "'Carga sem soma dos eixos via integracao Shopee' (Modern) — preso na fila ha 458h.",
    "'Cargas travadas - Aguardando emissao' — assunto repetido (6x).",
    "'Geracao de Carga Magazine Luiza' com 211h para resolver.",
    "'Emissao Vrid' (Amazon) com quase 300h para resolver.",
 ],
 'causa_raiz': (
    "Cada marketplace tem seu proprio modelo de lastro (carga_meli, arquivo_amazon, arquivo_shopee) e pontos "
    "de falha proprios — ex.: modo de falha documentado 'estacao de origem nao cadastrada -> sem CIOT' no "
    "Shopee Linehaul 3PL (#17601), com mecanismo exato ainda em aberto na KB."
 ),
 'diagnostico': [
    "Identificar o marketplace e o dominio correspondente na KB.",
    "Verificar se a carga tem o lastro correto (ticket_carga_emiteai / arquivo correspondente).",
    "Para Shopee: checar cadastro de expedidor/estacao de origem antes de escalar como bug.",
 ],
 'solucao': [
    "Escalar bugs estruturais direto para o time de dominio, evitando reprocessamento manual repetido sem causa raiz.",
    "Usar o runbook do #17601 (KB shopee/ciot) como referencia para CIOT ausente em Linehaul 3PL.",
 ],
 'prevencao': [
    "Dashboard de saude por marketplace (cargas travadas, tempo medio ate emissao).",
    "Fechar a lacuna do #17601 (mecanismo exato do gate de CIOT no Shopee Linehaul).",
 ],
 'confianca': 'Media — volume confirmado pelos chamados; causas especificas por marketplace nem todas investigadas em codigo.',
})

topics.append({
 'slug': '05_cte_pendente',
 'titulo': 'CT-e pendente para emissao / rejeitado',
 'subtitulo': 'Diagnostico de CT-e travado e seu efeito cascata em CIOT/MDF-e',
 'ref': "cluster 'CTE PENDENTE PARA EMISSAO' (14x) e 'Cte rejeitado' (13x)",
 'descricao': (
    "94 chamados no mes, com destaque para 'CTE PENDENTE PARA EMISSAO' (14x) e 'Cte rejeitado' (13x). "
    "Tempo medio de 45,7h, entre os mais altos do levantamento."
 ),
 'metricas': [
    ('Total de chamados CT-e', '94'),
    ('Tempo medio de resolucao', '45,7 horas'),
    ('Chamados com mais de 48h em aberto', '30'),
 ],
 'sintomas': [
    "CT-e nao sai do estado pendente por longos periodos.",
    "Rejeicao recorrente sem causa clara para o operador.",
    "'Erro na emissao do CTe e nao geracao de CIOT e MDFe' (efeito cascata).",
 ],
 'causa_raiz': (
    "CT-e pendente/rejeitado costuma estar ligado a falha no pipeline coreografado (montagem/XML/assinatura/"
    "envio) via Event Control, ou a documento referenciado cancelado / inconsistencia de IE. Como CIOT e MDF-e "
    "dependem do CT-e, uma falha aqui se propaga."
 ),
 'diagnostico': [
    "Checar Event Control / Ticket para ver em qual etapa do pipeline o CT-e travou.",
    "Verificar se ha documento referenciado (ex.: CT-e complementar) cancelado.",
    "Conferir inscricao estadual (IE) e cadastro tributario da empresa envolvida.",
 ],
 'solucao': [
    "Reprocessar pelo Event Control quando a causa for transiente.",
    "Escalar para o dominio quando a rejeicao persistir apos reprocessamento.",
 ],
 'prevencao': [
    "Alerta automatico para CT-e pendente ha mais de N horas.",
    "FAQ com as causas mais comuns de rejeicao de CT-e.",
 ],
 'confianca': 'Media — mecanismo geral documentado na KB do dominio cte; sem leitura profunda do L3 nesta sessao.',
})

topics.append({
 'slug': '06_performance_instabilidade',
 'titulo': 'Performance, instabilidade e morosidade do sistema',
 'subtitulo': 'O cluster com o pior tempo de resolucao do mes',
 'ref': "'MOROSIDADE NO SISTEMA' - 334 horas para resolver",
 'descricao': (
    "78 chamados de lentidao/instabilidade, com o pior tempo de resolucao do levantamento: 'MOROSIDADE NO "
    "SISTEMA' levou 314h (~13 dias) e 'instabilidade no site' 277h."
 ),
 'metricas': [
    ('Total de chamados de performance/instabilidade', '78'),
    ('Tempo medio de resolucao', '36,7 horas'),
    ('Pior caso do mes', "'MOROSIDADE NO SISTEMA' - 334 horas"),
    ('Chamados de "Monitoramento Proativo" no mes', 'apenas 6'),
 ],
 'sintomas': [
    "'MOROSIDADE NO SISTEMA' e 'Morosidade' (7 + 6 ocorrencias).",
    "'instabilidade no site' — alta urgencia, 277h para resolver.",
    "Sistema lento sem incidente formal aberto pela infra antes da reclamacao do cliente.",
 ],
 'causa_raiz': (
    "Nao investigada em profundidade nesta sessao. O padrao — tempos muito acima da media dos demais clusters — "
    "sugere que esses chamados dependem de escalonamento manual para infra/plataforma em vez de deteccao automatica."
 ),
 'diagnostico': [
    "Verificar correlacao de horario com picos de carga (fim de dia, fechamento de mes).",
    "Checar filas RabbitMQ e escala KEDA no momento do chamado.",
    "Verificar bloqueios (locks) ou queries lentas no Postgres do tenant no horario relatado.",
 ],
 'solucao': [
    "Acionar diretamente o time de plataforma/infra, nao tratar como chamado funcional generico.",
    "Registrar horario exato e tenant afetado para correlacionar com dashboards de infraestrutura.",
 ],
 'prevencao': [
    "Ampliar o uso da categoria 'Monitoramento Proativo' (hoje subutilizada).",
    "Alerta automatico de latencia/fila represada que abra chamado interno antes do cliente perceber.",
    "SLA diferenciado para chamados de instabilidade.",
 ],
 'confianca': 'Baixa para causa raiz tecnica (fora do escopo desta analise); alta para os numeros do Movidesk.',
})

topics.append({
 'slug': '07_mdfe_ciot',
 'titulo': 'MDF-e acoplado ao CIOT (emissao / manifesto)',
 'subtitulo': 'Quando o problema de MDF-e e, na pratica, um problema de CIOT',
 'ref': "cluster 'Problema de Emissao - Manifesto - CIOT' (15 ocorrencias)",
 'descricao': (
    "75 chamados no mes, com destaque para 'Problema de Emissao - Manifesto - CIOT' (15 ocorrencias), "
    "evidenciando que boa parte dos problemas de MDF-e sao, na pratica, problemas de CIOT que bloqueiam o manifesto."
 ),
 'metricas': [
    ('Total de chamados MDF-e', '75'),
    ('Tempo medio de resolucao', '40 horas'),
    ('Cluster mais recorrente', "'Problema de Emissao - Manifesto - CIOT' (15x)"),
 ],
 'sintomas': [
    "MDF-e nao emite porque o CIOT nao foi gerado/autorizado (flag 'bloquear MDF-e sem CIOT' ativa).",
    "'NAO ESTA GERANDO MDFe NA EMISSAO DE DOCUMENTOS' — 334h para resolver.",
    "'Emissao de MDFe MELI - Validacao CIOT MLP' — 459h na fila de bugs.",
 ],
 'causa_raiz': (
    "Quando a carga e de lotacao com a flag 'bloquear MDF-e sem CIOT' ativa, o MDF-e so emite apos o CIOT "
    "existir. Se o CIOT falhar silenciosamente (nenhuma parametrizacao casa), o MDF-e fica bloqueado sem "
    "mensagem clara para o operador."
 ),
 'diagnostico': [
    "Verificar se existe CIOT para a carga — se nao existe, investigar por que (ver manual 02).",
    "Confirmar se a flag 'bloquear MDF-e sem CIOT' esta ativa para o transportador/tomador.",
    "Verificar UF intermediaria / Cadastro de Trajetos se a rejeicao for 'trajeto invalido'.",
 ],
 'solucao': [
    "Resolver a causa raiz no CIOT antes de tentar reemitir o MDF-e.",
    "Se o problema for cadastro de trajeto, corrigir e reprocessar.",
 ],
 'prevencao': [
    "Mensagem de erro mais clara no MDF-e quando bloqueado por ausencia de CIOT.",
    "Dashboard cruzando MDF-e bloqueado x CIOT ausente.",
 ],
 'confianca': 'Media-alta — mecanismo documentado na KB (dominio mdfe e ciot); casos especificos (MELI/MLP) nao investigados em codigo.',
})

topics.append({
 'slug': '08_ciot_preso_processando',
 'titulo': 'CIOT preso em "processando" ou ausente no relatorio',
 'subtitulo': 'Dois mecanismos distintos, dois tratamentos distintos',
 'ref': "'CIOT PROCESSANDO NAO APARECE NO RELATORIO DE CIOTS' - 413h em aberto",
 'descricao': (
    "Padrao especifico dentro do cluster CIOT: CIOT que nunca sai do estado PROCESSANDO e nem aparece na "
    "Consulta de CIOT. Ha chamado ativo no mes com exatamente esse sintoma."
 ),
 'metricas': [
    ('Chamado de referencia', "'CIOT PROCESSANDO NAO APARECE NO RELATORIO DE CIOTS' (413h)"),
    ('Fixes documentados na KB', '#17960 (rejeicao cadastral) e #17346 (race Persiste-vs-Envia)'),
 ],
 'sintomas': [
    "CIOT nunca sai de 'processando'.",
    "CIOT nem aparece na tela de Consulta de CIOT (numero_ciot nulo).",
    "Reprocessar pelo Event Control as vezes resolve, as vezes nao.",
 ],
 'causa_raiz': (
    "Dois mecanismos distintos geram o mesmo sintoma: (a) race de concorrencia entre consumers Persiste e "
    "Envia (#17346) — reprocessar resolve; (b) rejeicao cadastral (placa/veiculo/transportador nao cadastrado) "
    "que, antes do #17960, deixava o CIOT preso permanentemente sem nunca ter sido criado no provedor — "
    "reprocessar NAO resolve nesse caso."
 ),
 'diagnostico': [
    "Primeiro passo obrigatorio: verificar se o cadastro de placa/veiculo/transportador esta completo.",
    "Se o cadastro esta correto, tratar como poll travado — reprocessar via Event Control.",
    "Se o cadastro estava incompleto, regularizar ANTES de qualquer reprocessamento.",
 ],
 'solucao': [
    "Separar os dois caminhos no atendimento para nao perder tempo reprocessando um caso cadastral.",
    "Confirmar apos a correcao se o CIOT passou a INCONSISTENTE ou AUTORIZADO.",
 ],
 'prevencao': [
    "Runbook com arvore de decisao: cadastro incompleto -> corrigir cadastro; cadastro ok -> reprocessar.",
    "Validar se o fix #17960 cobre 100% dos fluxos — o chamado ativo sugere caminho nao coberto.",
 ],
 'confianca': 'Alta — mecanismo duplo documentado no L2/L3 da KB (dominio ciot), incluindo PRs #17346 e #17960.',
})

topics.append({
 'slug': '09_nfse',
 'titulo': 'NFS-e: demora e erro de emissao',
 'subtitulo': 'Baixo volume, mas o maior tempo medio de resolucao do mes',
 'ref': "Tempo medio de resolucao: 80,9 horas (o mais alto do levantamento)",
 'descricao': (
    "Apenas 15 chamados no mes, mas com o MAIOR tempo medio de resolucao de todo o levantamento: 80,9 horas. "
    "Indica atendimento mais lento e/ou menos padronizado, mesmo com baixo volume."
 ),
 'metricas': [
    ('Total de chamados NFS-e', '15'),
    ('Tempo medio de resolucao', '80,9 horas (o mais alto do levantamento)'),
    ('Chamados com mais de 48h em aberto', '7 (quase metade do volume total)'),
 ],
 'sintomas': [
    "'ERRO EMISSAO DE NFSe' parado na fila de bugs ha mais de 450h.",
    "Emissao presa em 'Processando' ou retorno de 'Negacao Sistemica' sem tratativa clara.",
 ],
 'causa_raiz': (
    "A NFS-e tem TRES caminhos de emissao diferentes conforme o municipio (gateway e-Notas / emissor proprio "
    "para Sao Paulo e Belo Horizonte / servico nacional dedicado). O suporte pode nao ter clareza imediata de "
    "qual caminho o municipio do cliente usa, o que explica o tempo medio muito mais alto que os demais dominios."
 ),
 'diagnostico': [
    "Identificar o municipio (codigo IBGE) da empresa e qual dos 3 caminhos se aplica.",
    "Para Sao Paulo (3550308) e Belo Horizonte (3106200): checar o emissor proprio no monolito backend.",
    "Para os demais: verificar o gateway e-Notas ou o servico nacional dedicado.",
    "Confirmar certificado digital e inscricao municipal antes de investigar como bug.",
 ],
 'solucao': [
    "Direcionar o chamado para a fila certa desde a abertura, conforme o caminho de emissao identificado.",
 ],
 'prevencao': [
    "Guia rapido de triagem NFS-e por municipio.",
    "Campo/tag no Movidesk para o caminho de emissao (e-Notas / proprio SP-BH / nacional) na abertura do chamado.",
 ],
 'confianca': 'Media-alta para a estrutura dos 3 caminhos (KB); a causa exata da demora e inferencia a partir do padrao de tempos.',
})

topics.append({
 'slug': '10_cadastro_importacao',
 'titulo': 'Cadastro (motorista / veiculo / importacao) com erro',
 'subtitulo': 'Armadilhas conhecidas de parsing de CSV e falha silenciosa no proprietario',
 'ref': "Tempo medio de resolucao: 69 horas (segundo maior do levantamento)",
 'descricao': (
    "Volume baixo (17 chamados) mas segundo maior tempo medio de resolucao: 69 horas. Envolve principalmente "
    "erros de importacao em massa (CSV) de motorista/veiculo/proprietario."
 ),
 'metricas': [
    ('Total de chamados de cadastro', '17'),
    ('Tempo medio de resolucao', '69 horas'),
    ('Chamados presos em fila de bugs', '3'),
 ],
 'sintomas': [
    "Erro ao importar CSV de motoristas/veiculos.",
    "'erro de alfa numerico no cadastro de ocorrencias' — problema de validacao de campo.",
    "Falha silenciosa no proprietario durante upload em massa.",
 ],
 'causa_raiz': (
    "Armadilhas conhecidas de parsing no CSV: o cabecalho usa virgula mas os dados usam ponto-e-virgula (ou "
    "vice-versa) em alguns templates, e ha um 'silent-drop' conhecido no cadastro de proprietario durante a "
    "importacao em massa — o registro falha sem erro visivel ao usuario."
 ),
 'diagnostico': [
    "Verificar o delimitador do CSV enviado pelo cliente contra o padrao esperado pelo template.",
    "Para proprietario 'desaparecido' pos-importacao: checar se houve silent-drop.",
    "Confirmar vinculo motorista-exclusivo e FKs de proprietario/beneficiario quando o erro for de vinculo.",
 ],
 'solucao': [
    "Orientar o cliente a exportar o CSV no delimitador correto antes de reenviar.",
    "Para silent-drop confirmado, tratar como bug de dominio (frota) e escalar.",
 ],
 'prevencao': [
    "Validacao explicita de delimitador na tela de upload, com mensagem de erro clara.",
    "Eliminar o silent-drop do proprietario — devolver erro visivel em vez de descartar silenciosamente.",
    "FAQ/template padrao de CSV disponibilizado ao cliente antes da importacao.",
 ],
 'confianca': 'Alta para as armadilhas de parsing (documentadas explicitamente na KB do dominio frota).',
})

# ============================================================
# GERACAO DOS ARQUIVOS
# ============================================================
generated = []
for i, sec in enumerate(topics, 1):
    doc, section = new_document()
    add_cover(doc, sec['titulo'], sec['subtitulo'], sec.get('ref'))
    add_topic_body(doc, sec)
    import re as _re
    safe_title = _re.sub(r'[\\/:*?"<>|]', '', sec['titulo'])
    fname = f"{i:02d} - {safe_title}.docx"
    path = os.path.join(OUT_DIR, fname)
    doc.save(path)
    generated.append(path)
    print("OK:", path)

print(f"\nTotal gerado: {len(generated)} manuais em {OUT_DIR}")
