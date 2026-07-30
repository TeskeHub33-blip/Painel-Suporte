import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
ACCENT = RGBColor(0x2E, 0x5C, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual de Suporte")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = DARK
    for _ in range(2):
        doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("10 Situacoes Recorrentes - CIOT / VPO / MDF-e / CT-e / NFS-e / Cadastro / Performance")
    r2.font.size = Pt(14)
    r2.font.color.rgb = ACCENT
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Baseado na analise de 1.099 chamados do Movidesk (julho/2026) e na base de conhecimento tecnica (horus-mcp / knowledge-base)")
    r3.italic = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = GREY
    doc.add_page_break()

def add_toc_placeholder(topics):
    h = doc.add_heading("Indice", level=1)
    for i, t in enumerate(topics, 1):
        p = doc.add_paragraph(f"{i}. {t['titulo']}", style='List Number')
    doc.add_page_break()

def add_meta_table(rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = 'Indicador'
    hdr[1].text = 'Valor (julho/2026)'
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    doc.add_paragraph()

def add_bullets(items, style='List Bullet'):
    for it in items:
        doc.add_paragraph(it, style=style)

def add_section(sec):
    doc.add_heading(sec['titulo'], level=1)

    doc.add_heading('Descricao', level=2)
    doc.add_paragraph(sec['descricao'])

    doc.add_heading('Indicadores no periodo analisado', level=2)
    add_meta_table(sec['metricas'])

    doc.add_heading('Sintomas tipicos (exemplos reais de chamados)', level=2)
    add_bullets(sec['sintomas'])

    doc.add_heading('Causa raiz', level=2)
    doc.add_paragraph(sec['causa_raiz'])

    doc.add_heading('Diagnostico passo a passo', level=2)
    add_bullets(sec['diagnostico'], style='List Number')

    doc.add_heading('Solucao / encaminhamento', level=2)
    add_bullets(sec['solucao'])

    doc.add_heading('Prevencao e monitoria preventiva sugerida', level=2)
    add_bullets(sec['prevencao'])

    if sec.get('confianca'):
        p = doc.add_paragraph()
        r = p.add_run('Nivel de confianca desta secao: ' + sec['confianca'])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    doc.add_page_break()

topics = []

# 1 --------------------------------------------------------------
topics.append({
 'titulo': '1. Reprocessar VPO (Vale-Pedagio) travado ou rejeitado',
 'descricao': (
    "E o assunto de maior volume no mes: variacoes de 'Reprocessar vpo', 'Vpo rejeitado', "
    "'Erro no VPO' somam mais de 70 chamados. Na maioria dos casos o vale-pedagio da viagem "
    "nao e comprado/emitido automaticamente e o operador de suporte precisa acionar reprocessamento manual "
    "pelo Event Control."
 ),
 'metricas': [
    ('Total de chamados (assunto relacionado a VPO/pedagio)', '176'),
    ('Tempo medio de resolucao', '19,6 horas'),
    ('Chamados com mais de 48h em aberto', '26'),
    ('Maior recorrencia de assunto identico', "'Reprocessar vpo' (42x) + 'Reprocessar vpo' variante (28x)"),
 ],
 'sintomas': [
    "Vale-pedagio nao aparece na viagem mesmo com CIOT autorizado.",
    "Chamado aberto com titulo generico 'Reprocessar vpo' / 'Vpo reprocessamento' / 'Erro no VPO'.",
    "Vale-pedagio com valor incorreto (chamado especifico visto no mes).",
    "Caso MELI: 'MELI - REPROCESSAMENTO DE VPO' e 'ERRO VPO - MELI' aparecem como cluster proprio (14 + 7 ocorrencias) — sugere um problema mais especifico na integracao Mercado Livre.",
 ],
 'causa_raiz': (
    "Segundo a KB (dominio vale-pedagio), o VPO depende do CIOT estar autorizado — a compra do vale-pedagio "
    "e disparada pelo evento CIOT_AUTORIZADO, e o identificador do pedagio vem gravado no payload do CIOT "
    "sem evento proprio (nao ha outbox transacional CIOT->VPO). Isso cria uma janela de corrida (race) onde "
    "o consumer do VPO pode nao encontrar o identificador a tempo, resultando em VPO nao emitido/orfao. "
    "Tambem existe o caminho 'VPO sem CIOT' (proveniencia carga-sem-ciot, via REPOM VPO) que tem regras proprias "
    "e pode falhar separadamente."
 ),
 'diagnostico': [
    "Confirmar o status do CIOT da viagem (deve estar AUTORIZADO — sem isso o VPO nunca sera comprado).",
    "Verificar no Event Control / Ticket a timeline do documento (ticket_uuid) buscando o evento de compra de VPO.",
    "Checar se e um caso 'VPO orfao' (CIOT autorizado mas identificador de pedagio nao propagado a tempo).",
    "Para casos MELI: verificar se a rota/HU tem lastro correto (ticket_carga_emiteai) antes de reprocessar.",
    "Se o CIOT esta rejeitado/cancelado, o VPO nao deve ser reprocessado sem antes resolver o CIOT.",
 ],
 'solucao': [
    "Reprocessar o evento de compra de VPO pelo Event Control apos confirmar CIOT autorizado.",
    "Se for 'VPO orfao' recorrente, escalar para o time de dominio vale-pedagio com o ticket_uuid + timestamp da corrida — este e um padrao de bug conhecido (falta de outbox transacional), nao apenas operacional.",
    "Para o cluster MELI, tratar como fila separada — pode exigir fix na integracao, nao apenas reprocessamento.",
 ],
 'prevencao': [
    "Criar alerta automatico quando um CIOT fica AUTORIZADO ha mais de X minutos sem VPO associado (detecta o caso orfao antes do cliente abrir chamado).",
    "Adicionar ao runbook de suporte um passo obrigatorio: checar status do CIOT antes de qualquer reprocessamento de VPO.",
    "Levar para o backlog de produto a criacao do outbox transacional CIOT->VPO (fechar a lacuna estrutural documentada na KB).",
    "FAQ dedicado para o cluster MELI-VPO, com passo a passo especifico dessa integracao.",
 ],
 'confianca': 'Alta para o mecanismo geral (documentado na KB do dominio vale-pedagio); media para a causa exata do cluster MELI (nao investigada em codigo nesta sessao).',
})

# 2 --------------------------------------------------------------
topics.append({
 'titulo': '2. Adiantamento do CIOT nao reflete no EmiteAi (fluxo REPOM/provedor)',
 'descricao': (
    "Cliente informa/gera o CIOT com um percentual de adiantamento acordado com o provedor (ex.: REPOM), "
    "mas esse valor nao aparece no EmiteAi. Caso real identificado no mes: chamado #29077 "
    "'REPOM - Adiantamento no CIOT - Esta refletindo na Repom e nao na Emiteai em componentes + CIOTS em "
    "contigencia nao retornam para Emiteai' (aberto ha 427h no momento da analise, no status 'Aguardando Time CS')."
 ),
 'metricas': [
    ('Chamados no cluster CIOT', '139'),
    ('Tempo medio de resolucao (CIOT)', '47,5 horas — o mais alto entre os top 5 clusters por volume'),
    ('Chamados de CIOT com mais de 48h em aberto', '39'),
    ('Chamado de referencia', '#29077 (Bug, Media urgencia, aberto ha 427h)'),
 ],
 'sintomas': [
    "Percentual de adiantamento (ex.: 70%) negociado com o provedor nao aparece em nenhum relatorio/tela do EmiteAi.",
    "'VALOR CIOT VINDO COM O VALOR ERRADO' (chamado de alta urgencia, aberto ha 413h).",
    "CIOTs emitidos em contingencia (codigo verificador 'XXXX') nao retornam/atualizam no EmiteAi.",
 ],
 'causa_raiz': (
    "Existem duas telas distintas para tratar CIOT manualmente: (1) 'Gestao de terceiros > Consulta de CIOTs > + Novo' "
    "(emissao manual/avulso), que TEM o campo '(%) de adiantamento' e grava em ciot.valor_adiantamento; e "
    "(2) 'Informar CIOT' (Emissor / conferencia Shopee-Amazon), que apenas vincula um numero de CIOT ja existente "
    "e NAO possui campo de adiantamento. Quando o cliente gera o CIOT direto no provedor (com adiantamento la parametrizado) "
    "e apenas informa o numero no EmiteAi, o valor nunca e capturado — nao e bug, e limitacao de fluxo. "
    "Adicionalmente, a parametrizacao de CIOT tem o campo 'percentual_adiantamento_fornecedor' que, se nulo, "
    "faz o calculo automatico do fornecedor tambem nao aplicar nenhum adiantamento."
 ),
 'diagnostico': [
    "Identificar qual tela/fluxo o cliente usa para o CIOT: emissao manual (+Novo) ou 'Informar CIOT'.",
    "Consultar a tabela parametrizacao_ciot do tenant e verificar percentual_adiantamento_fornecedor (nulo = nao configurado).",
    "Verificar tambem 'criar_conta_pagar_geracao_ciot' e 'gerar_lancamento_adiantamento_separado' — se false, nao havera lancamento financeiro mesmo com percentual configurado.",
    "Para casos de contingencia: verificar se o codigo_verificador esta gravado como 'XXXX' (sentinela) — isso e esperado em autorizacao por contingencia, nao e erro; o valor real chega depois via nova consulta ao provedor.",
 ],
 'solucao': [
    "Se o cliente usa 'Informar CIOT': orientar a migrar para a emissao manual (+Novo) quando quiser que o EmiteAi capture o adiantamento, OU configurar percentual_adiantamento_fornecedor na parametrizacao para calculo automatico via carga.",
    "Ativar criar_conta_pagar_geracao_ciot e gerar_lancamento_adiantamento_separado se o cliente espera ver o lancamento financeiro.",
    "Para contingencia: orientar que 'XXXX' e temporario e que o sistema deve reconsultar o provedor — se ficar preso, tratar como caso 8 (CIOT preso em processando) deste manual.",
 ],
 'prevencao': [
    "Adicionar validacao/alerta na tela de parametrizacao de CIOT: se percentual_adiantamento_fornecedor estiver nulo, exibir aviso ao cadastrar.",
    "FAQ objetivo: 'Por que o adiantamento do meu CIOT nao aparece no EmiteAi?' explicando a diferenca entre emitir e informar CIOT.",
    "Monitoria: alertar quando um tenant usa majoritariamente 'Informar CIOT' e tem reclamacoes de adiantamento — pode ser sinal de expectativa desalinhada com o produto.",
 ],
 'confianca': 'Alta — mecanismo confirmado por consulta direta as tabelas ciot e parametrizacao_ciot em producao (tenant JM Line) e cruzado com o chamado real #29077.',
})

# 3 --------------------------------------------------------------
topics.append({
 'titulo': '3. CIOT rejeitado / cancelamento / encerramento manual',
 'descricao': (
    "Chamados pedindo para cancelar, encerrar ou tratar rejeicao de CIOT manualmente. Inclui variacoes "
    "'Cte/Ciot rejeitado' (21 ocorrencias combinadas), 'cancelar ciot' (4x) e pedidos pontuais como "
    "'JRPS - Encerrar CIOTs' e 'Encerramentos CIOTs' (Rodoind)."
 ),
 'metricas': [
    ('Chamados de CIOT rejeitado/cancelamento (subset do cluster CIOT)', '~25-30'),
    ('Tempo medio do cluster CIOT', '47,5 horas'),
    ('Chamados presos em fila de Bugs relacionados', '11'),
 ],
 'sintomas': [
    "CIOT aparece como REJEITADO mesmo com a viagem em transito no provedor.",
    "Pedido para 'Cancelar CIOT' ou 'Encerrar CIOT' feito manualmente pelo suporte (ex.: 'Encerrar CIOT - 520010806545').",
    "Cliente pede reemissao apos rejeicao repetida.",
 ],
 'causa_raiz': (
    "A KB do dominio CIOT documenta que 'ja cadastrado' (ErrorCode 35 na REPOM) e uma resposta de idempotencia, "
    "nao uma rejeicao real — se o suporte trata isso como rejeicao, o CIOT valido e cancelado/reemitido "
    "incorretamente. Rejeicao cadastral verdadeira (placa/veiculo/transportador nao cadastrado) e outro caminho, "
    "corrigido no #17960 para marcar como INCONSISTENTE (nao mais preso em PROCESSANDO)."
 ),
 'diagnostico': [
    "Antes de cancelar/encerrar manualmente, checar no provedor (REPOM/etc.) se o CIOT esta realmente ativo — se sim, NAO tratar como rejeitado.",
    "Se rejeicao e por cadastro (placa/veiculo/transportador ausente na base), regularizar o cadastro primeiro — reprocessar nao resolve nesse caso.",
    "Para 'ja cadastrado', confiar no registro local autorizado, sem gerar novo CIOT.",
 ],
 'solucao': [
    "Padronizar um checklist de suporte: 1) checar estado real no provedor, 2) so entao decidir cancelar/reemitir.",
    "Encaminhar para dev quando o cadastro ja esta correto e o CIOT segue rejeitado sem motivo aparente.",
 ],
 'prevencao': [
    "Treinamento do time de suporte no comportamento de idempotencia do CIOT (ErrorCode 35) para evitar cancelamentos indevidos.",
    "Checklist fixo no Movidesk (campo obrigatorio) antes de fechar chamados desse tipo.",
 ],
 'confianca': 'Alta — casos e correcoes documentados na KB (L2/L3 do dominio ciot).',
})

# 4 --------------------------------------------------------------
topics.append({
 'titulo': '4. Integracoes de marketplace (Amazon / Shopee / Meli / Magalu)',
 'descricao': (
    "120 chamados no mes envolvendo as integracoes de marketplace, com tempo medio de 43,1h — o segundo maior "
    "entre os clusters de alto volume. O subconjunto mais recorrente e VPO/CIOT dentro do fluxo Meli (ja coberto "
    "no item 1), mas ha tambem problemas estruturais de carga (eixos, integracao Shopee, cargas travadas)."
 ),
 'metricas': [
    ('Total de chamados marketplace', '120'),
    ('Tempo medio de resolucao', '43,1 horas'),
    ('Chamados com mais de 48h em aberto', '35'),
 ],
 'sintomas': [
    "'Carga sem soma dos eixos quando recebemos via integracao Shopee' (Modern) — bug preso na fila ha 458h.",
    "'Cargas travadas - Aguardando emissao' — assunto repetido (6x).",
    "'Geracao de Carga Magazine Luiza' com 211h para resolver.",
    "'Emissao Vrid' (Amazon) com quase 300h para resolver.",
 ],
 'causa_raiz': (
    "Segundo a KB, cada marketplace tem seu proprio modelo de lastro (carga_meli, arquivo_amazon, arquivo_shopee) "
    "e pontos de falha proprios — por exemplo, o modo de falha documentado 'estacao de origem nao cadastrada -> sem CIOT' "
    "no Shopee Linehaul 3PL (#17601), com mecanismo exato ainda em aberto na KB."
 ),
 'diagnostico': [
    "Identificar o marketplace e o dominio correspondente (meli / amazon / shopee / ecommerce na KB).",
    "Verificar se a carga tem o lastro correto (ticket_carga_emiteai / arquivo correspondente).",
    "Para Shopee especificamente: checar cadastro de expedidor/estacao de origem antes de escalar como bug.",
 ],
 'solucao': [
    "Escalar bugs estruturais (ex.: soma de eixos, VRID travado) direto para o time de dominio, evitando reprocessamento manual repetido sem causa raiz.",
    "Usar o runbook de investigacao do #17601 (KB dominio shopee/ciot) como referencia para casos de CIOT ausente em Linehaul 3PL.",
 ],
 'prevencao': [
    "Dashboard de saude por marketplace (cargas travadas, tempo medio ate emissao) para deteccao proativa.",
    "Fechar a lacuna do #17601 (mecanismo exato do gate de geracao de CIOT no Shopee Linehaul) — hoje e uma pergunta em aberto na KB.",
 ],
 'confianca': 'Media — volume e padroes confirmados pelos chamados; causas especificas variam por marketplace e nem todas foram investigadas em codigo nesta sessao.',
})

# 5 --------------------------------------------------------------
topics.append({
 'titulo': '5. CT-e pendente para emissao / rejeitado',
 'descricao': (
    "94 chamados no mes, com destaque para 'CTE PENDENTE PARA EMISSAO' (14x) e 'Cte rejeitado' (13x). "
    "Tempo medio de resolucao de 45,7h, entre os mais altos do levantamento."
 ),
 'metricas': [
    ('Total de chamados CT-e', '94'),
    ('Tempo medio de resolucao', '45,7 horas'),
    ('Chamados com mais de 48h em aberto', '30'),
 ],
 'sintomas': [
    "CT-e nao sai do estado pendente por longos periodos.",
    "Rejeicao recorrente sem causa clara para o operador.",
    "'Demora na emissao de CTE' e 'Erro na emissao do CTe e nao geracao de CIOT e MDFe' (efeito cascata).",
 ],
 'causa_raiz': (
    "Pela KB do dominio cte, CT-e pendente/rejeitado costuma estar ligado a Event Control (falha no pipeline "
    "coreografado montagem/XML/assinatura/envio) ou a documento referenciado cancelado / inconsistencia de IE. "
    "Como o CIOT e o MDF-e dependem do CT-e, uma falha aqui se propaga (efeito cascata visto em varios chamados)."
 ),
 'diagnostico': [
    "Checar Event Control / Ticket para ver em qual etapa do pipeline (montagem/XML/assinatura/envio) o CT-e travou.",
    "Verificar se ha documento referenciado (ex.: CT-e complementar) cancelado, o que bloqueia o fluxo.",
    "Conferir inscricao estadual (IE) e cadastro tributario da empresa envolvida.",
 ],
 'solucao': [
    "Reprocessar pelo Event Control quando a causa for transiente (mesma logica usada para CIOT preso).",
    "Escalar para o dominio quando a rejeicao persistir apos reprocessamento — pode ser regra tributaria ou cadastral.",
 ],
 'prevencao': [
    "Alerta automatico para CT-e pendente ha mais de N horas, antes que o cliente precise abrir chamado.",
    "FAQ com as causas mais comuns de rejeicao de CT-e e o que o operador pode verificar antes de escalar.",
 ],
 'confianca': 'Media — mecanismo geral documentado na KB do dominio cte; nao foi feita leitura profunda do L3 nesta sessao.',
})

# 6 --------------------------------------------------------------
topics.append({
 'titulo': '6. Performance / instabilidade / morosidade do sistema',
 'descricao': (
    "78 chamados de lentidao/instabilidade, com o pior tempo de resolucao do levantamento: 'MOROSIDADE NO SISTEMA' "
    "levou 314 horas (~13 dias) e 'instabilidade no site' 277 horas."
 ),
 'metricas': [
    ('Total de chamados de performance/instabilidade', '78'),
    ('Tempo medio de resolucao', '36,7 horas'),
    ('Pior caso do mes', "'MOROSIDADE NO SISTEMA' — 334 horas"),
    ('Categoria "Monitoramento Proativo" usada no mes', 'apenas 6 chamados'),
 ],
 'sintomas': [
    "'MOROSIDADE NO SISTEMA' e 'Morosidade' (7 + 6 ocorrencias).",
    "'instabilidade no site' — chamado de alta urgencia, 277h para resolver.",
    "Sistema lento sem incidente formal aberto pela area de infra antes da reclamacao do cliente.",
 ],
 'causa_raiz': (
    "Nao investigada em profundidade nesta sessao (exigiria acesso a metricas de infraestrutura/APM). "
    "O padrao observado — tempos de resolucao muito acima da media do restante dos clusters — sugere que "
    "esses chamados dependem de escalonamento manual para a area de infra/plataforma em vez de deteccao automatica."
 ),
 'diagnostico': [
    "Verificar se ha correlacao de horario com picos de carga (ex.: fim de dia, fechamento de mes).",
    "Checar filas RabbitMQ (mensageria) e escala KEDA no momento do chamado — fila represada e sintoma comum de lentidao.",
    "Verificar se o Postgres do tenant apresenta bloqueios (locks) ou queries lentas concorrentes no horario relatado.",
 ],
 'solucao': [
    "Acionar diretamente o time de plataforma/infra ao inves de tratar como chamado funcional generico.",
    "Registrar o horario exato e o tenant afetado para correlacionar com dashboards de infraestrutura.",
 ],
 'prevencao': [
    "Ampliar o uso da categoria 'Monitoramento Proativo' (hoje so 6 chamados no mes) — sinaliza que a deteccao proativa de degradacao ainda e subutilizada.",
    "Implementar alerta automatico de latencia/fila represada que abra chamado interno ANTES do cliente perceber.",
    "SLA diferenciado para chamados de instabilidade, dado o tempo de resolucao muito acima da media.",
 ],
 'confianca': 'Baixa para causa raiz tecnica (fora do escopo desta analise, que foi baseada em ticket + KB de dominio, sem acesso a APM/infra); alta para os numeros do Movidesk.',
})

# 7 --------------------------------------------------------------
topics.append({
 'titulo': '7. MDF-e acoplado ao CIOT (emissao/manifesto)',
 'descricao': (
    "75 chamados no mes, com destaque para o cluster 'Problema de Emissao - Manifesto - CIOT' (15 ocorrencias), "
    "evidenciando que boa parte dos problemas de MDF-e sao, na pratica, problemas de CIOT que bloqueiam o manifesto."
 ),
 'metricas': [
    ('Total de chamados MDF-e', '75'),
    ('Tempo medio de resolucao', '40 horas'),
    ('Cluster mais recorrente', "'Problema de Emissao - Manifesto - CIOT' (15x)"),
 ],
 'sintomas': [
    "MDF-e nao emite porque o CIOT nao foi gerado/autorizado (flag 'bloquear MDF-e sem CIOT' ativa).",
    "'NAO ESTA GERANDO MDFe NA EMISSAO DE DOCUMENTOS' — bug de alto impacto, 334h para resolver.",
    "'Emissao de MDFe MELI - Validacao CIOT MLP' — parado ha 459h na fila de bugs.",
 ],
 'causa_raiz': (
    "A KB do dominio mdfe documenta que, quando a carga e de lotacao com a flag 'bloquear MDF-e sem CIOT' ativa, "
    "o MDF-e so emite apos o CIOT existir. Se o CIOT falhar silenciosamente (ver item 2 do dominio ciot: falha "
    "silenciosa quando nenhuma parametrizacao casa), o MDF-e fica bloqueado sem mensagem clara para o operador."
 ),
 'diagnostico': [
    "Verificar primeiro se existe CIOT para a carga — se nao existe, investigar por que ele nao foi gerado (ver item 2 deste manual e a secao 'CIOT nunca nasce' da KB).",
    "Confirmar se a flag 'bloquear MDF-e sem CIOT' esta ativa para o transportador/tomador.",
    "Verificar UF intermediaria / Cadastro de Trajetos se a rejeicao for 'trajeto invalido'.",
 ],
 'solucao': [
    "Resolver a causa raiz no CIOT antes de tentar reemitir o MDF-e — reemitir o MDF-e sem CIOT valido nao resolve.",
    "Se o problema for cadastro de trajeto, corrigir e reprocessar.",
 ],
 'prevencao': [
    "Mensagem de erro mais clara no MDF-e quando o bloqueio for por ausencia de CIOT (hoje o sintoma aparece so a jusante).",
    "Dashboard cruzando MDF-e bloqueado x CIOT ausente para identificar o padrao antes da abertura de chamado.",
 ],
 'confianca': 'Media-alta — mecanismo de bloqueio documentado na KB (dominio mdfe e ciot); casos especificos (ex.: MELI/MLP) nao investigados em codigo.',
})

# 8 --------------------------------------------------------------
topics.append({
 'titulo': '8. CIOT preso em "processando" ou ausente no relatorio',
 'descricao': (
    "Padrao especifico dentro do cluster CIOT: CIOT que nunca sai do estado PROCESSANDO e, por isso, nem "
    "aparece na tela de Consulta de CIOT. Ha um chamado ativo no mes exatamente com esse sintoma: "
    "'CIOT PROCESSANDO NAO APARECE NO RELATORIO DE CIOTS' (aberto ha 413h)."
 ),
 'metricas': [
    ('Chamado de referencia no mes', "'CIOT PROCESSANDO NAO APARECE NO RELATORIO DE CIOTS' (413h em aberto)"),
    ('Fix documentado na KB para este padrao', '#17960 (rejeicao cadastral) e #17346 (race Persiste-vs-Envia)'),
 ],
 'sintomas': [
    "CIOT nunca sai de 'processando', independente de quanto tempo passe.",
    "O CIOT nem aparece na tela de Consulta de CIOT (numero_ciot fica nulo).",
    "Reprocessar pelo Event Control as vezes resolve, as vezes nao (dois mecanismos diferentes por tras do mesmo sintoma).",
 ],
 'causa_raiz': (
    "A KB documenta DOIS mecanismos distintos que geram o mesmo sintoma: (a) poll travado / race de concorrencia "
    "entre os consumers Persiste e Envia (#17346) — nesse caso reprocessar pelo Event Control resolve; "
    "e (b) rejeicao cadastral (placa/veiculo/transportador nao cadastrado) que, antes do #17960, escapava do "
    "tratamento de erro e deixava o CIOT preso permanentemente SEM nunca ter sido criado no provedor — nesse "
    "caso reprocessar NAO resolve, pois nao ha CIOT valido para consultar."
 ),
 'diagnostico': [
    "Primeiro passo obrigatorio: verificar se o cadastro de placa/veiculo/transportador esta completo e correto.",
    "Se o cadastro esta correto, tratar como poll travado — reprocessar via Event Control (ticket + chave de acesso).",
    "Se o cadastro estava incompleto, regularizar o cadastro ANTES de qualquer tentativa de reprocessamento (reprocessar nao resolve o caso cadastral).",
 ],
 'solucao': [
    "Separar claramente os dois caminhos no atendimento (poll travado vs rejeicao cadastral) para nao perder tempo reprocessando um caso que exige correcao de cadastro.",
    "Confirmar apos a correcao se o CIOT passou a status INCONSISTENTE (esperado pos-#17960) ou AUTORIZADO.",
 ],
 'prevencao': [
    "Runbook de suporte com arvore de decisao: cadastro incompleto? -> corrigir cadastro. Cadastro ok? -> reprocessar Event Control.",
    "Validar se o fix #17960 esta de fato cobrindo 100% dos fluxos — o chamado ativo no mes sugere que pode haver um caminho nao coberto (investigar em codigo antes de fechar o chamado como duplicata do fix antigo).",
 ],
 'confianca': 'Alta — mecanismo duplo documentado em detalhe no L2/L3 da KB do dominio ciot, incluindo os PRs de correcao (#17346, #17960).',
})

# 9 --------------------------------------------------------------
topics.append({
 'titulo': '9. NFS-e — demora e erro de emissao',
 'descricao': (
    "Apenas 15 chamados no mes, mas com o MAIOR tempo medio de resolucao de todo o levantamento: 80,9 horas. "
    "Isso indica atendimento mais lento e/ou menos padronizado para esse dominio especifico, mesmo com baixo volume."
 ),
 'metricas': [
    ('Total de chamados NFS-e', '15'),
    ('Tempo medio de resolucao', '80,9 horas (o mais alto do levantamento)'),
    ('Chamados com mais de 48h em aberto', '7 (quase metade do volume total)'),
 ],
 'sintomas': [
    "'ERRO EMISSAO DE NFSe' parado na fila de bugs ha mais de 450h.",
    "Emissao presa em 'Processando' ou retorno de 'Negacao Sistemica' sem tratativa clara.",
 ],
 'causa_raiz': (
    "Pela KB, a NFS-e tem TRES caminhos de emissao diferentes conforme o municipio (gateway e-Notas / emissor "
    "proprio para Sao Paulo e Belo Horizonte / servico nacional dedicado), cada um com seu proprio gate e "
    "particularidades. O suporte pode nao ter clareza imediata de qual caminho o municipio do cliente usa, "
    "o que explica o tempo medio de resolucao muito mais alto que os demais dominios."
 ),
 'diagnostico': [
    "Identificar o municipio (codigo IBGE) da empresa e, por ele, qual dos 3 caminhos de emissao se aplica.",
    "Para Sao Paulo (3550308) e Belo Horizonte (3106200): checar o emissor proprio no monolito backend.",
    "Para os demais: verificar o gateway e-Notas (caminho padrao/fallback) ou o servico nacional dedicado.",
    "Confirmar certificado digital e inscricao municipal da empresa antes de investigar como bug.",
 ],
 'solucao': [
    "Direcionar o chamado para a fila certa desde a abertura, de acordo com o caminho de emissao identificado — hoje parece nao haver esse roteamento, o que explica a demora.",
 ],
 'prevencao': [
    "Criar um guia rapido de triagem NFS-e por municipio (qual caminho usar, o que checar primeiro) para reduzir o tempo medio de 80,9h.",
    "Adicionar campo/tag no Movidesk para o caminho de emissao (e-Notas / proprio SP-BH / nacional) logo na abertura do chamado.",
 ],
 'confianca': 'Media-alta para a estrutura dos 3 caminhos (documentada na KB); a causa exata da demora no suporte e uma inferencia a partir do padrao de tempos, nao confirmada com o time de CS.',
})

# 10 --------------------------------------------------------------
topics.append({
 'titulo': '10. Cadastro (motorista / veiculo / importacao) com erro',
 'descricao': (
    "Volume baixo (17 chamados) mas segundo maior tempo medio de resolucao do levantamento: 69 horas. "
    "Envolve principalmente erros de importacao em massa (CSV) de motorista/veiculo/proprietario."
 ),
 'metricas': [
    ('Total de chamados de cadastro', '17'),
    ('Tempo medio de resolucao', '69 horas'),
    ('Chamados presos em fila de bugs', '3'),
 ],
 'sintomas': [
    "Erro ao importar CSV de motoristas/veiculos ('Importacao de motoristas e veiculos').",
    "'Nidos - erro de alfa numerico no cadastro de ocorrencias' — problema de validacao de campo.",
    "Falha silenciosa no proprietario durante upload em massa.",
 ],
 'causa_raiz': (
    "A KB do dominio frota documenta armadilhas conhecidas de parsing no CSV: o cabecalho usa virgula mas os "
    "dados usam ponto-e-virgula (ou vice-versa) em alguns templates, e ha um 'silent-drop' conhecido no "
    "cadastro de proprietario durante a importacao em massa — ou seja, o registro falha sem erro visivel ao usuario."
 ),
 'diagnostico': [
    "Verificar o delimitador do arquivo CSV enviado pelo cliente (virgula vs ponto-e-virgula) contra o padrao esperado pelo template.",
    "Para casos de proprietario 'desaparecido' pos-importacao: checar se houve silent-drop (o registro nao gerou erro mas tambem nao foi criado).",
    "Confirmar vinculo motorista-exclusivo e FKs de proprietario/beneficiario quando o erro for de vinculo, nao de parsing.",
 ],
 'solucao': [
    "Orientar o cliente a exportar o CSV no delimitador correto antes de reenviar.",
    "Para silent-drop confirmado, tratar como bug de dominio (frota) e escalar, pois hoje o sistema nao avisa o usuario.",
 ],
 'prevencao': [
    "Validacao explicita de delimitador na tela de upload, com mensagem de erro clara (elimina a maior causa de retrabalho).",
    "Eliminar o silent-drop do proprietario — devolver erro visivel em vez de descartar silenciosamente.",
    "FAQ/template padrao de CSV disponibilizado ao cliente antes da importacao, reduzindo erros na origem.",
 ],
 'confianca': 'Alta para as armadilhas de parsing (documentadas explicitamente na KB do dominio frota).',
})

add_title_page()
add_toc_placeholder(topics)

intro = doc.add_heading('Sobre este manual', level=1)
doc.add_paragraph(
    "Este documento consolida as 10 situacoes de maior impacto identificadas na analise dos chamados do "
    "Movidesk de julho/2026 (1.099 chamados completos, sem amostragem), cruzadas com a base de conhecimento "
    "tecnica (dominio-a-dominio) do EmiteAi. Cada secao segue a mesma estrutura: descricao do problema, "
    "indicadores do periodo, sintomas tipicos com exemplos reais, causa raiz (quando documentada), diagnostico "
    "passo a passo, solucao/encaminhamento e sugestoes de monitoria preventiva."
)
doc.add_paragraph(
    "O nivel de confianca de cada secao esta indicado ao final — secoes 'alta confianca' foram cruzadas "
    "diretamente com a KB tecnica (L1/L2/L3) ou com consultas diretas ao banco de producao; secoes de "
    "'media/baixa confianca' sao inferidas a partir do padrao dos chamados e merecem validacao adicional "
    "com o time de dominio antes de virarem processo oficial."
)
doc.add_page_break()

for t in topics:
    add_section(t)

out_path = "C:/Users/WellingtonErvinoTesk/Documents/Claude/manual_suporte/Manual_Suporte_10_Situacoes.docx"
doc.save(out_path)
print("Salvo em:", out_path)
